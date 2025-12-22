from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path
from typing import Any, Dict, List

import yaml
from docx import Document

from src.config import AppConfig


# =============================
# Data models
# =============================

@dataclass
class Company:
    name: str
    location: str = ""
    website: str = ""
    careers: str = ""
    contact_email: str = ""
    stack_hints: List[str] = None
    domain: str = ""
    why: str = ""
    notes: str = ""


@dataclass
class Profile:
    person: Dict[str, Any]
    education: Dict[str, Any]
    profile: Dict[str, Any]


# =============================
# Loaders
# =============================

def load_companies(path: str = "companies.yaml") -> List[Company]:
    raw = yaml.safe_load(Path(path).read_text(encoding="utf-8")) or {}
    items = raw.get("companies", []) or []

    out: List[Company] = []
    for c in items:
        out.append(
            Company(
                name=c.get("name", "") or "",
                location=c.get("location", "") or "",
                website=c.get("website", "") or "",
                careers=c.get("careers", "") or "",
                contact_email=c.get("contact_email", "") or "",
                stack_hints=c.get("stack_hints", []) or [],
                domain=c.get("domain", "") or "",
                why=c.get("why", "") or "",
                notes=c.get("notes", "") or "",
            )
        )
    return out


def load_profile(path: str = "profile.yaml") -> Profile:
    raw = yaml.safe_load(Path(path).read_text(encoding="utf-8")) or {}
    return Profile(
        person=raw.get("person", {}) or {},
        education=raw.get("education", {}) or {},
        profile=raw.get("profile", {}) or {},
    )


def slugify(name: str) -> str:
    keep = []
    for ch in name.strip():
        if ch.isalnum() or ch in (" ", "_", "-"):
            keep.append(ch)
    return "".join(keep).strip().replace(" ", "_")[:60]


# =============================
# Helpers
# =============================

def _join(items: List[str]) -> str:
    items = [str(x).strip() for x in (items or []) if str(x).strip()]
    return ", ".join(items)


def _lia_start(cfg: AppConfig, p: Profile) -> str:
    # Prefer profile.yaml; fall back to config
    return p.education.get("lia_target_start") or getattr(cfg.lia.target, "desired_start", "oktober 2026")


def _company_stack(company: Company) -> str:
    return _join(company.stack_hints) if company.stack_hints else "Java/Spring och modern webbutveckling"


def _contact_block(p: Profile) -> str:
    person = p.person
    bits = []
    for k in ("phone", "email", "linkedin", "github"):
        v = person.get(k)
        if v:
            bits.append(str(v))
    return "  |  ".join(bits)


def company_alignment_paragraph_sv(company: Company) -> str:
    """
    Company-specific paragraph from companies.yaml:
    - why (1 sentence)
    - domain (label)
    - notes (free text)
    """
    parts: List[str] = []
    if company.why:
        parts.append(company.why.strip())
    if company.domain:
        parts.append(f"Branschfokus: {company.domain.strip()}.")
    if company.notes:
        parts.append(company.notes.strip())
    return " ".join([p for p in parts if p]).strip()


# =============================
# Outreach email (SV) with mode toggle
# =============================

def draft_email_sv(cfg: AppConfig, company: Company, p: Profile, mode: str = "cold") -> tuple[str, str]:
    """
    mode:
      - "cold"         = proactive outreach (no job ad)
      - "application"  = responding to an ad / role
    """
    full_name = p.person.get("full_name", "")
    program = p.education.get("program", "Javautvecklare")
    school = p.education.get("school", "Nackademin")
    lia_start = _lia_start(cfg, p)
    stack = _company_stack(company)

    if mode == "application":
        subject = f"Ansökan: LIA (start {lia_start}) – Fullstack (Java + frontend)"
        intro = (
            f"Jag vill anmäla intresse för en LIA-plats hos {company.name} med start {lia_start} (HT26). "
            f"Jag studerar till {program} på {school} och söker en roll med fullstackfokus."
        )
    else:
        subject = f"LIA (start {lia_start}) – Fullstack (Java backend + frontend) – Förfrågan"
        intro = (
            f"Jag heter {full_name} och studerar till {program} på {school}. "
            f"Jag söker en LIA-plats med start {lia_start} (HT26) och hör av mig proaktivt."
        )

    align = company_alignment_paragraph_sv(company)

    body = (
        f"Hej!\n\n"
        f"{intro}\n\n"
        f"Jag är intresserad av {company.name} eftersom ni arbetar med {stack}. "
        f"Jag trivs i miljöer där kvalitet, lärande och samarbete är centralt, och vill gärna bidra praktiskt i teamet "
        f"– både i Java-backend och i frontend.\n\n"
        f"{align}\n\n"
        f"Under utbildningen har jag byggt praktiska projekt i Java (bl.a. Best Gym Ever och Greenest Hotel Plant Program) "
        f"och arbetar strukturerat med clean code, TDD och Git. Utöver utbildningen har jag byggt en Telegram-bot för nyheter "
        f"med AI-sammanfattningar samt mitt eget verktyg för att bevaka och strukturera LIA-spår.\n\n"
        f"Skulle ni kunna tänka er att ta emot en LIA-student under hösten 2026? "
        f"Jag skickar gärna CV (PDF) och personligt brev, eller bokar ett kort samtal.\n\n"
        f"Vänliga hälsningar,\n"
        f"{full_name}\n"
        f"{_contact_block(p)}\n"
    )
    return subject, body


def write_outreach_email_txt(folder: Path, subject: str, body: str) -> Path:
    folder.mkdir(parents=True, exist_ok=True)
    path = folder / "outreach_email.txt"
    path.write_text(f"SUBJECT: {subject}\n\n{body}", encoding="utf-8")
    return path


# =============================
# LinkedIn DM
# =============================

def draft_linkedin_dm_sv(cfg: AppConfig, company: Company, p: Profile) -> str:
    full_name = p.person.get("full_name", "")
    lia_start = _lia_start(cfg, p)
    stack = _company_stack(company)

    return (
        f"Hej! Jag heter {full_name} och studerar till Javautvecklare på Nackademin. "
        f"Jag söker LIA med start {lia_start} (HT26) och har fullstackfokus (Java backend + frontend). "
        f"Jag såg att ni jobbar med {stack} och undrar om ni kan tänka er att ta emot en LIA-student hösten 2026? "
        f"Jag skickar gärna CV + kort presentation. Tack!"
    )


def write_linkedin_dm_txt(folder: Path, message: str) -> Path:
    folder.mkdir(parents=True, exist_ok=True)
    path = folder / "linkedin_dm.txt"
    path.write_text(message, encoding="utf-8")
    return path


# =============================
# DOCX generation
# =============================

def write_personligt_brev_docx(
    folder: Path,
    cfg: AppConfig,
    company: Company,
    p: Profile,
    variant: str,
) -> Path:
    """
    variant: "kort" or "standard"
    """
    folder.mkdir(parents=True, exist_ok=True)

    full_name = p.person.get("full_name", "")
    program = p.education.get("program", "Javautvecklare")
    school = p.education.get("school", "Nackademin")
    lia_start = _lia_start(cfg, p)
    stack = _company_stack(company)

    backend = _join(p.profile.get("backend_strengths", []))
    frontend = _join(p.profile.get("frontend_strengths", []))
    prior = (p.profile.get("prior_experience_summary", "") or "").strip()

    doc = Document()
    doc.add_heading("Personligt brev", level=1)

    # Company-specific alignment early (but after doc exists)
    align = company_alignment_paragraph_sv(company)
    if align:
        doc.add_paragraph(align)

    doc.add_paragraph(f"Hej {company.name},")
    doc.add_paragraph(
        f"Jag studerar till {program} på {school} och söker en LIA-plats med start {lia_start} (HT26). "
        f"Jag är intresserad av {company.name} eftersom er inriktning matchar det jag vill utvecklas inom: {stack}."
    )

    if variant == "kort":
        doc.add_paragraph(
            f"Tekniskt trivs jag i backend med {backend} och vill även bidra i frontend med {frontend}. "
            f"Jag uppskattar arbetssätt med code reviews, tydlig struktur och fokus på kvalitet."
        )
    else:
        doc.add_paragraph(
            f"Under utbildningen har jag byggt flera praktiska projekt i Java och arbetat med ett tydligt fokus på "
            f"struktur, testbar kod och versionering. Tekniskt trivs jag i backend med {backend}. "
            f"Jag vill samtidigt bidra och utvecklas i frontend med {frontend} och bygga features end-to-end – "
            f"från API till gränssnitt – med fokus på kvalitet och begriplig kod."
        )

        doc.add_paragraph(
            "Som exempel har jag byggt:\n"
            "• Best Gym Ever och Greenest Hotel Plant Program (Java-projekt med OOP, struktur och clean code)\n"
            "• Telegram News Bot (SV→RU) med AI-sammanfattningar\n"
            "• LIA Finder AI Assistant (verktyg för bevakning och outreach kring LIA/praktik)\n"
        )

        if prior:
            doc.add_paragraph(prior)

    doc.add_paragraph(
        "Jag bifogar mitt CV (PDF) och berättar gärna mer i ett kort samtal. Tack för att ni tar er tid att läsa."
    )
    doc.add_paragraph("Vänliga hälsningar,")
    doc.add_paragraph(full_name)
    doc.add_paragraph(_contact_block(p))

    filename = "personligt_brev_kort.docx" if variant == "kort" else "personligt_brev_standard.docx"
    path = folder / filename
    doc.save(str(path))
    return path


def write_cv_highlights_docx(folder: Path, cfg: AppConfig, company: Company, p: Profile) -> Path:
    folder.mkdir(parents=True, exist_ok=True)

    full_name = p.person.get("full_name", "")
    lia_start = _lia_start(cfg, p)
    stack = _company_stack(company)

    backend = _join(p.profile.get("backend_strengths", []))
    frontend = _join(p.profile.get("frontend_strengths", []))

    doc = Document()
    doc.add_heading("CV Highlights – LIA Fullstack (Java + Frontend)", level=1)
    doc.add_paragraph(f"Namn: {full_name}")
    doc.add_paragraph(f"Mål: LIA start {lia_start} (HT26)")
    doc.add_paragraph(f"Företag: {company.name}")
    doc.add_paragraph(f"Fokus: {stack}")

    doc.add_heading("Styrkor", level=2)
    doc.add_paragraph(f"• Backend: {backend}")
    doc.add_paragraph(f"• Frontend: {frontend}")
    doc.add_paragraph("• Arbetssätt: Git, struktur, ansvarstagande, samarbete och kontinuerligt lärande.")

    path = folder / "cv_highlights.docx"
    doc.save(str(path))
    return path
